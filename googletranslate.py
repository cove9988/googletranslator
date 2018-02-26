from docx import Document
from googletrans import Translator
import datetime

'''
destination = 
    'af': 'afrikaans',
    'sq': 'albanian',
    'am': 'amharic',
    'ar': 'arabic',
    'hy': 'armenian',
    'az': 'azerbaijani',
    'eu': 'basque',
    'be': 'belarusian',
    'bn': 'bengali',
    'bs': 'bosnian',
    'bg': 'bulgarian',
    'ca': 'catalan',
    'ceb': 'cebuano',
    'ny': 'chichewa',
    'zh-cn': 'chinese (simplified)',
    'zh-tw': 'chinese (traditional)',
    'co': 'corsican',
    'hr': 'croatian',
    'cs': 'czech',
    'da': 'danish',
    'nl': 'dutch',
    'en': 'english',
    'eo': 'esperanto',
    'et': 'estonian',
    'tl': 'filipino',
    'fi': 'finnish',
    'fr': 'french',
    'fy': 'frisian',
    'gl': 'galician',
    'ka': 'georgian',
    'de': 'german',
    'el': 'greek',
    'gu': 'gujarati',
    'ht': 'haitian creole',
    'ha': 'hausa',
    'haw': 'hawaiian',
    'iw': 'hebrew',
    'hi': 'hindi',
    'hmn': 'hmong',
    'hu': 'hungarian',
    'is': 'icelandic',
    'ig': 'igbo',
    'id': 'indonesian',
    'ga': 'irish',
    'it': 'italian',
    'ja': 'japanese',
    'jw': 'javanese',
    'kn': 'kannada',
    'kk': 'kazakh',
    'km': 'khmer',
    'ko': 'korean',
    'ku': 'kurdish (kurmanji)',
    'ky': 'kyrgyz',
    'lo': 'lao',
    'la': 'latin',
    'lv': 'latvian',
    'lt': 'lithuanian',
    'lb': 'luxembourgish',
    'mk': 'macedonian',
    'mg': 'malagasy',
    'ms': 'malay',
    'ml': 'malayalam',
    'mt': 'maltese',
    'mi': 'maori',
    'mr': 'marathi',
    'mn': 'mongolian',
    'my': 'myanmar (burmese)',
    'ne': 'nepali',
    'no': 'norwegian',
    'ps': 'pashto',
    'fa': 'persian',
    'pl': 'polish',
    'pt': 'portuguese',
    'pa': 'punjabi',
    'ro': 'romanian',
    'ru': 'russian',
    'sm': 'samoan',
    'gd': 'scots gaelic',
    'sr': 'serbian',
    'st': 'sesotho',
    'sn': 'shona',
    'sd': 'sindhi',
    'si': 'sinhala',
    'sk': 'slovak',
    'sl': 'slovenian',
    'so': 'somali',
    'es': 'spanish',
    'su': 'sundanese',
    'sw': 'swahili',
    'sv': 'swedish',
    'tg': 'tajik',
    'ta': 'tamil',
    'te': 'telugu',
    'th': 'thai',
    'tr': 'turkish',
    'uk': 'ukrainian',
    'ur': 'urdu',
    'uz': 'uzbek',
    'vi': 'vietnamese',
    'cy': 'welsh',
    'xh': 'xhosa',
    'yi': 'yiddish',
    'yo': 'yoruba',
    'zu': 'zulu',
    'fil': 'Filipino',
    'he': 'Hebrew'
mix = True, will create a doc with oiginial and translated language line by line.
      False, will only create a doc with translated language.  
'''
def work_doc(filename,destination = 'zh-CN', mix = True):
    tx = lambda t : Translator().translate(t,dest=destination).text
    dlt = datetime.datetime.now()
    print('start to translate ....{0}...to...{1}'.format(filename,destination))
    doc = Document(filename)
    for cnt, p in enumerate(doc.paragraphs):
        if len(p.text) > 0 :
            try:        
                p.text = p.text + ('\n' + tx(p.text) if mix else '')
            except:
                print('Error: ',p.text)
    print('translated {0} paragraphs in {1} secs'.format(cnt,(datetime.datetime.now() - dlt).total_seconds()))
    dlt = datetime.datetime.now()
    for cnt,table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                if len(cell.text) > 0:
                    try:
                        cell.text = cell.text + ('\n' + tx(cell.text) if mix else '')
                    except :
                        print('Error: ',cell.text )

    print('translated {0} paragraphs in {1} secs'.format(cnt,(datetime.datetime.now() - dlt).total_seconds()))
    f = filename.replace('.doc',destination + '.doc')
    doc.save(f)
    print('done, save the tranlated file as: {0}'.format(f))
if __name__ == '__main__':
    filename = 'P2.docx'
    work_doc(filename)
